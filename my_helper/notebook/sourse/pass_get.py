from PyQt5 import uic
from PyQt5.QtWidgets import QDialog, QMessageBox
from PyQt5.QtCore import Qt
from my_helper.notebook.sourse.inserts import get_from_db
import datetime as dt
import os
import docx
import docxtpl
#  сделать мессаджбоксы на Сохранить
count_days = (31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31)
main_file = "B:/my_helper/getpass.docx"
print_file = "B:/my_helper/to_print/getpass.docx"
designer_file = '../designer_ui/pass_get.ui'


class GetPass(QDialog):
    def __init__(self, parent):
        super(GetPass, self).__init__()
        uic.loadUi(designer_file, self)
        # pass
        self.parent = parent
        self.b_ok.clicked.connect(self.ev_ok)
        self.b_cancel.clicked.connect(self.ev_cancel)
        self.b_open.clicked.connect(self.my_open_file)

        self.d_note.setDate(dt.datetime.now().date())
        self.d_from.setDate(dt.datetime.now().date())
        self.d_to.setDate(dt.datetime.now().date())
        self.number.setValue(self.parent.get_next_number())

        self.list_ui = (self.worker_1, self.worker_2, self.worker_3, self.worker_4, self.worker_5,
                        self.worker_6, self.worker_7, self.worker_8, self.worker_9, self.worker_10)
        self.data = {"customer": "", "company": "", "start_date": "", "end_date": "",
                     "contract": "", "date_contract": "", "number": "", "data": ""}
        if not self.init_workers() and self.init_contracts():
            self.close()

    def init_contracts(self):
        self.cb_contracts.addItem("(нет)")
        contracts = self.parent.db.get_data("name, date", "contract")
        if not contracts:
            return False
        for row in contracts:
            self.cb_contract.addItem(row[0])
        return True

    def init_workers(self):
        for item in self.list_ui:
            item.addItem("(нет)")
            item.activated[str].connect(self.new_worker)
            item.setEnabled(False)
        self.list_ui[0].setEnabled(True)
        people = self.parent.db.get_data("family, name, surname, post, passport, "
                                            "passport_got, birthday, adr,  live_adr", "workers")
        if not people:
            return False
        for name in people:
            family = name[0] + " " + ".".join([name[1][0], name[2][0]]) + "."
            for item in self.list_ui:
                item.addItem(family)
        return True

    # получить данные
    # для заполнения текста
    def get_data(self):
        self.data["start_date"] = self.d_from.text()
        self.data["end_date"] = self.d_to.text()
        self.data["customer"] = self.parent.customer
        self.data["company"] = self.parent.company
        for contract in self.from_db("number, date", "contract"):
            if self.cb_contract.currentText() == contract[0]:
                self.data["contract"] = contract[0]
                self.data["date_contract"] = contract[1]

    def get_worker(self, family):
        rows = self.from_db("family, name, surname, post, passport, "
                            "passport_got, birthday, adr,  live_adr", "workers")
        if family == "all":
            return rows
        for row in rows:
            if family[:-5] == row[0]:  # на форме фамилия в виде Фамилия И.О.
                return row

    # обработчики кнопок
    def ev_ok(self):
        self.data["number"] = "Исх. № " + self.number.text()
        self.data["date"] = "от. " + self.d_note.text()

        self.get_data()
        doc = docxtpl.DocxTemplate(main_file)
        doc.render(self.data)

        doc.save(print_file)
        doc = docx.Document(print_file)
        # Заполнить таблицу
        workers = []
        for elem in self.list_ui:
            if elem.currentText() != "(нет)":
                workers.append(self.get_worker(elem.currentText()))
        i = 1
        for people in workers:
            doc.tables[1].add_row()
            doc.tables[1].rows[i].cells[0].text = str(i)
            doc.tables[1].rows[i].cells[1].text = " ".join(people[0:2])
            doc.tables[1].rows[i].cells[2].text = people[3]
            doc.tables[1].rows[i].cells[3].text = people[6]
            doc.tables[1].rows[i].cells[4].text = " ".join(people[4:6])
            doc.tables[1].rows[i].cells[5].text = people[7]
            doc.tables[1].rows[i].cells[6].text = people[8]
            i += 1
        doc.save(print_file)
        self.close()
        os.startfile(print_file)

    def new_worker(self):
        flag = True
        for item in self.list_ui:
            if item.currentText() != "(нет)":
                item.setEnabled(True)
            else:
                item.setEnabled(flag)
                flag = False
        pass

    def ev_cancel(self):
        self.close()

    def save_pattern(self):
        data = {"who": self.cb_who.text(),
                "object_name": self.cb_object.text(),
                "workers": self.get_list()}
        self.zip_pattern(data)
        # запоковать в словарь
        # сохранить в файл

    def my_open_file(self):
        print("open file")
        pass

    def kill_pattern(self):
        pass

    def from_db(self, fields, table):
        self.parent.db.execute(get_from_db(fields, table))
        return self.parent.db.fetchall()




