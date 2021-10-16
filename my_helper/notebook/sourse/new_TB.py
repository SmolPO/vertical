
from PyQt5.QtWidgets import QMessageBox as mes
from PyQt5.QtWidgets import QDialog, QCheckBox, QLabel
from PyQt5 import uic
import docx
import docxtpl
import os
designer_file = '../designer_ui/new_TB.ui'
_types = {"1": "/OT.docx", "2": "/PTM.docx", "3": "/ES.docx"}


class NewTB(QDialog):
    def __init__(self, parent=None):
        super(NewTB, self).__init__()
        # pass
        uic.loadUi(designer_file, self)
        self.parent = parent
        self.table = "workers"
        self.count = self.parent.count_people_tb
        self.b_ok.clicked.connect(self.ev_ok)
        self.b_cancel.clicked.connect(self.ev_cancel)
        self.rows_from_db = self.parent.db.get_data("*", self.table)
        self.main_file = ""
        self.print_file = ""
        self.list_ui = list()
        self.init_list()

    def init_list(self):
        """
        создать таблицу 0 | ФИО
        добавить все ФИО работающих сотрудников
        :return:
        """
        g = iter(range(self.count + 1))
        for item in self.rows_from_db:
            i = next(g)
            number = QLabel(item[-1])
            checkbox = QCheckBox()
            family = QLabel(item[0] + item[1][0] + "." + item[2][0] + ".")
            self.grid.addWidget(number, i, 0)
            self.grid.addWidget(checkbox, i, 1)
            self.grid.addWidget(family, i, 2)
        return

    def get_list_people(self):
        list_people = list()
        for i in range(self.count):
            if self.grid.itemAtPosition(i, 1).widget().isChecked():
                number = self.grid.itemAtPosition(i, 0).widget().text()
                list_people.append(self.rows_from_db[int(number)-1])
        return list_people

    def ev_ok(self):
        people = self.get_list_people()
        self.create_protocols(people)
        return True

    def ev_cancel(self):
        self.close()

    def create_protocols(self, people):
        """
        список протоколов: просмотреть всех сотрудников и найти номера протоколов
        :param people: список id сотрудников
        :return:
        """
        printed_docs = list()
        worker = list()
        fields = [0, 1, 2, 4, 20, 21, 22, 24]
        _key = 20
        for man in people:
            if not man[_key] in printed_docs:
                for field in fields:
                    worker.append(man[field])
                print(worker)
                for key in _types.keys():
                    self.print_doc(worker, key)
                printed_docs.append(man[_key])
        return

    def print_doc(self, workers, number_type):
        data = dict()
        data["date"] = workers[7] + "/" + str(number_type)  # номер столбца с датой
        data["number_doc"] = workers[6]  # номер протокол
        self.create_table(workers, number_type)
        doc = docxtpl.DocxTemplate(self.main_file + _types[number_type])
        doc.render(self.data)
        try:
            doc.save(self.print_file + _types[number_type])
        except:
            self.close()

        os.startfile(self.print_file)
        return

    def create_table(self, data, number_type):
        g = iter(range(len(data)))
        print(data)
        doc = docx.Document(self.print_file + _types[str(number_type)])
        for item in data:
            i = next(g)
            doc.tables[1].add_row()
            doc.tables[1].rows[i].cells[0].text = str(i)
            doc.tables[1].rows[i].cells[1].text = " ".join(item[:3]) # ФИО
            doc.tables[1].rows[i].cells[2].text = item[3]   # профессия
            doc.tables[1].rows[i].cells[3].text = "Сдал №" + item[5]
        doc.save(self.print_file + _types[str(number_type)])


class CountPeople(QDialog):
    def __init__(self, parent=None):
        super(CountPeople, self).__init__(parent)
        ui_file = '../designer_ui/count.ui'
        # pass
        uic.loadUi(ui_file, self)
        self.parent = parent
        self.table = "workers"
        self.b_ok.clicked.connect(self.ev_ok)
        self.b_cancel.clicked.connect(self.ev_cancel)
        count_people = len(self.parent.db.get_data("*", self.table))
        self.count.setMaximum(count_people)
        print(count_people)

    def ev_ok(self):
        self.parent.count_people_tb = self.count.value()
        self.close()
        return

    def ev_cancel(self):
        self.parent.count_people_tb = -1
        self.close()