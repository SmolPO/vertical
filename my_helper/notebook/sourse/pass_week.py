import docx
from pass_template import TempPass
from database import *


class WeekPass(TempPass):
    def __init__(self, parent):
        self.status_ = True
        self.conf = Ini(self)
        ui_file = self.conf.get_path_ui("pass_week")
        if not ui_file or ui_file == ERR:
            self.status_ = False
            return
        super(WeekPass, self).__init__(ui_file, parent, "contracts")
        self.parent = parent
        self.table = "contracts"
        if not self.status_:
            return
        self.rows_from_db = self.parent.db.get_data("*", self.table)
        self.d_from.setDate(dt.datetime.now().date())
        self.d_to.setDate(dt.datetime.now().date())
        self.workers = [self.worker_1, self.worker_2, self.worker_3,
                        self.worker_4, self.worker_5, self.worker_6,
                        self.worker_7, self.worker_8, self.worker_9]
        self.b_save.clicked.connect(self.save_pattern)
        self.b_kill.clicked.connect(self.kill_pattern)
        self.cb_other.stateChanged.connect(self.other_days)
        self.cb_sun.stateChanged.connect(self.week_days)
        self.cb_sub.stateChanged.connect(self.week_days)
        self.d_from.setEnabled(False)
        self.d_to.setEnabled(False)
        self.data = {"number": "", "date": "", "week_day": "", "contract": "", "type_work": "",
                     "part": "", "company": "", "customer": "", "post_boss": "", "boss_part": ""}
        self.list_ui = (self.worker_1, self.worker_2, self.worker_3, self.worker_4,
                        self.worker_5, self.worker_6, self.worker_7, self.worker_8, self.worker_9)
        if self.init_object() == ERR:
            self.status_ = False
            return
        if self.init_boss() == ERR:
            self.status_ = False
            return
        if self.init_workers() == ERR:
            self.status_ = False
            return
        self.main_file += "/pass_week.docx"

    # заполнение список
    def get_days(self):
        data = []
        now_weekday = dt.datetime.now().weekday()
        if self.cb_other.isChecked():
            data.append([self.d_from.text(), self.d_to.text()])
            return data
        if self.cb_sun.isChecked():
            sub_day = dt.datetime.now() + dt.timedelta(5 - now_weekday)
            data.append(".".join((str(sub_day.day), str(sub_day.month), str(sub_day.year))))
        if self.cb_sub.isChecked():
            sun_day = dt.datetime.now() + dt.timedelta(6 - now_weekday)
            data.append(".".join([str(sun_day.day), str(sun_day.month), str(sun_day.year)]))
        return data

    def init_object(self):
       for row in self.rows_from_db:
            self.cb_object.addItem(row[0])

    def init_boss(self):
        rows = self.parent.db.get_data("family, name, surname, post", "bosses")
        if rows == ERR:
            return msg_er(self, GET_DB)
        for people in rows:
            family = short_name(people[:3])
            self.cb_boss_part.addItem(family)       # брать из БД

    def init_workers(self):
        for item in self.list_ui:
            item.addItem(empty)
            item.activated[str].connect(self.new_worker)
            item.setEnabled(False)
        self.list_ui[0].setEnabled(True)
        for people in self.all_people:
            if people[-2] != 3:
                family = str(people[-1]) + ". " + short_name(people)
                for item in self.list_ui:
                    item.addItem(family)

    # обработчики кнопок
    def _ev_ok(self):
        if not self.cb_sun.isChecked() and not self.cb_sub.isChecked() and not self.cb_other.isChecked():
            mes.question(self, "Сообщение", "Укажите в какой день будете работать ", mes.Cancel)
            return False
        if self.cb_object.currentText() == empty:
            mes.question(self, "Сообщение", "Выберите объект", mes.Cancel)
            return False
        if self.cb_boss_part.currentText() == empty:
            mes.question(self, "Сообщение", "Выберите Босса", mes.Cancel)
            return False
        self.data["boss_part"] = self.cb_boss_part.currentText()
        self.data["post_boss"] = "Начальник цеха"
        if self.get_contract(self.cb_object.currentText()) == ERR:
            return
        if self.get_week_days() == ERR:
            return
        return True

        # Заполнить таблицу
    def _create_data(self, path):
        try:
            doc = docx.Document(path)
        except:
            msg_er(self, GET_FILE + path)
            return ERR
        i = 1
        list_people = list()
        for elem in self.list_ui:
            family = elem.currentText()
            if family != NOT:
                doc.tables[1].add_row()
                item = self.check_row(family)
                if item:
                    list_people.append(item)
        list_people.sort(key=lambda x: x[0])
        for people in list_people:
            doc.tables[1].rows[i].cells[0].text = str(i)
            doc.tables[1].rows[i].cells[1].text = " ".join(people[:3])
            doc.tables[1].rows[i].cells[2].text = people[3]
            doc.tables[1].rows[i].cells[3].text = people[6]
            doc.tables[1].rows[i].cells[4].text = " ".join(people[4:6])
            doc.tables[1].rows[i].cells[5].text = people[7]
            doc.tables[1].rows[i].cells[6].text = people[8]
            i += 1
        try:
            doc.save(path)
        except:
            return msg_er(self, GET_FILE + path)
        return True

    def check_row(self, row):
        my_id = int(row.split(".")[0])
        family = row.split(".")[1][1:-2]
        s_name = row.split(".")[1][-1]
        s_sur = row.split(".")[2]
        for item in self.all_people:
            if str(my_id) == str(item[-1]):
                if family == item[0]:
                    if s_name == item[1][0]:
                        if s_sur == item[2][0]:
                            return item

    def _get_data(self):
        return True

    def check_input(self):
        if self.list_ui[0].currentText() == empty:
            return msg_er(self, ADD_PEOPLE)
        return True
