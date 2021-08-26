import xml.etree.ElementTree as ET
import docx

def create_note(data):
    data = {"number": "1221", "data": "19.21.3221"}
    doc = docx.Document("B:/my_helper/week.docx")
    # номер исх
    doc.tables[0].rows[0].cells[0].text = "Исх. № " + data["number"]
    # дата
    doc.tables[0].rows[1].cells[0].text = "от. " + data["data"]
    # босс
    # заполняется вручную в шаблоне
    # Уважаемый вручную шаблоном
    # Просим Вас
    doc.paragraphs[6].add_run("Прошу Вас разрешить работы в выходной день 07.08.2021 по ремонту строительных конструкций (ограждающих стен и кровли галереи поз. Т4Б (825), пересыпок Б5, Б6) галерей конвейерных производства нитроаммофоски инв. 015409. цех по отгрузке минеральных удобрений Производства минеральных удобрений работникам ООО «Вертикаль» согласно договору 20/1-960-Р от 17.05.2021г с рабочей сменой с 07-00 до 19-00 часов:")
    # Заполнить таблицу
    doc.tables[1].rows[1].cells[0].text = "1"
    doc.tables[1].rows[1].cells[1].text = "Калентеенков Иван Сергеевич"
    doc.tables[1].rows[1].cells[2].text = "Прораб"
    doc.tables[1].rows[1].cells[3].text = "выдан"
    doc.tables[1].rows[1].cells[4].text = "рег"
    doc.tables[1].rows[1].cells[5].text = "прож"
    doc.tables[1].add_row()
    doc.tables[1].rows[2].cells[0].text = "1"
    doc.tables[1].rows[2].cells[1].text = "Калентеенков Иван Сергеевич"
    doc.tables[1].rows[2].cells[2].text = "Прораб"
    doc.tables[1].rows[2].cells[3].text = "выдан"
    doc.tables[1].rows[2].cells[4].text = "рег"
    doc.tables[1].rows[2].cells[5].text = "прож"
    doc.tables[1].add_row()

    doc.save("B:/my_helper/week.docx")

create_note("привет")

def create_file(data):
    root = ET.Element("xml")
    item_boss = ET.SubElement(root, "boss")
    item_boss.text = "boss"
    item_name = ET.SubElement(root, "name")
    item_name.text = "name"
    item_list = ET.SubElement(root, "list")
    for item in ["ads", "asdd", "wer"]:
        el = ET.SubElement(item_list, "worker")
        el.text = item
    tree = ET.ElementTree(root)
    tree.write("xml_test.xml")
# открыть или создать
try:
    tree = ET.parse("xml_test.xml")
    glob_root = tree.getroot()
    root = ET.SubElement(glob_root, "pattern")
    item_boss = ET.SubElement(root, "boss")
    item_boss.text = "boss"
    item_name = ET.SubElement(root, "name")
    item_name.text = "name"
    item_list = ET.SubElement(root, "list")
    for item in ["ads", "asdd", "wer"]:
        el = ET.SubElement(item_list, "worker")
        el.text = item
    tree.write("xml_test.xml")
except:
    create_file(["jgh1", "jbj2", "hj3"])
    pass



# добавить элемент
# найти элемент
# удалить элемент
