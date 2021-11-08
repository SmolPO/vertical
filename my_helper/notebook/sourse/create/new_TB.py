
from PyQt5.QtWidgets import QMessageBox as mes
from PyQt5.QtWidgets import QDialog, QCheckBox, QLabel
from PyQt5 import uic
import docx
import docxtpl
import os
from my_helper.notebook.sourse.database import *
types_docs = {"1": "/ot_doc.docx", "2": "/ptm_doc.docx", "3": "/eb_doc.docx"}
types_card = {"1": "/ot_card.docx", "2": "/ptm_card.docx", "3": "/es_card.docx"}


class NewTB(QDialog):
    def __init__(self, parent=None):
        self.status_ = True
        self.conf = Ini(self)
        ui_file = self.conf.get_path_ui("new_TB")
        if not ui_file:
            self.status_ = False
            return
        try:
            uic.loadUi(ui_file, self)
        except:
            self.status_ = False
            return
        super(NewTB, self).__init__()
        if not self.check_start():
            return
        self.parent = parent
        self.table = "workers"
        self.rows_from_db = self.parent.db.get_data("*", self.table)
        if self.rows_from_db == ERR:
            self.status_ = False
            return
        self.count = self.parent.count_people_tb
        self.b_ok.clicked.connect(self.ev_ok)
        self.b_cancel.clicked.connect(self.ev_cancel)
        self.path = dict()
        path_1 = self.conf.get_path("path")
        path_2 = self.conf.get_path("path_pat_tb")
        path_3 = self.conf.get_path("path_tb")
        if path_1 == ERR or path_2 == ERR or path_3 == ERR:
            self.status_ = False
            return
        self.path["main_folder"] = path_1 + path_2
        self.path["print_folder"] = path_1 + path_3
        self.list_ui = list()
        self.init_list()

    def init_list(self):
        g = iter(range(self.count + 1))
        for item in self.rows_from_db:
            i = next(g)
            number = QLabel(item[-1])
            checkbox = QCheckBox()
            family = QLabel(item[0] + item[1][0] + "." + item[2][0] + ".")
            self.grid.addWidget(number, i, 0)
            self.grid.addWidget(checkbox, i, 1)
            self.grid.addWidget(family, i, 2)
        return

    def get_list_people(self):
        list_people = list()
        for i in range(self.count):
            if self.grid.itemAtPosition(i, 1).widget().isChecked():
                number = self.grid.itemAtPosition(i, 0).widget().text()
                if len(self.rows_from_db) > int(number)-1:
                    list_people.append(self.rows_from_db[int(number)-1])
                else:
                    return msg_er(self, BIG_INDEX)
        return list_people

    def ev_ok(self):
        people = self.get_list_people()
        self.create_protocols(people)
        if self.create_cards(people) == ERR:
            return
        return True

    def ev_cancel(self):
        self.close()

    def create_protocols(self, people):
        printed_docs = list()
        worker = list()
        fields = [0, 1, 2, 4, 20, 21, 22, 24]
        _key = 20
        for man in people:
            if not man[_key] in printed_docs:
                for field in fields:
                    worker.append(man[field])
                print(worker)
                for key in types_docs.keys():
                    self.print_doc(worker, key)
                printed_docs.append(man[_key])
        return
    
    def create_cards(self, people):
        data = dict()
        for worker in people:
            data["family"] = " ".join(worker[:3])
            data["post"] = worker[3]
            data["number_doc"] = worker[20]
            data["number"] = worker[21]
            data["date"] = worker[22]
            for name in types_card.keys():
                try:
                    path = self.path["main_folder"] + types_card[name]
                    doc = docxtpl.DocxTemplate(path)
                    doc.render(self.data)
                    path = self.path["print_folder"] + types_card[name]
                    doc.save(path)
                    os.startfile(path, "print")
                except:
                    return msg_er(self, GET_FILE + path)
                self.close()

    def print_doc(self, workers, number_type):
        data = dict()
        data["date"] = workers[7] + "/" + str(number_type)  # номер столбца с датой
        data["number_doc"] = workers[6]  # номер протокол
        self.create_table(workers, number_type)
        try:
            path = self.path["main_folder"] + types_docs[number_type]
            doc = docxtpl.DocxTemplate(path)
            doc.render(self.data)
            path = self.path["print_folder"] + types_docs[number_type]
            doc.save(path)
            os.startfile(self.print_file)
        except:
            return msg_er(self, GET_FILE + path)

    def create_table(self, data, number_type):
        g = iter(range(len(data)))
        try:
            path = self.path["main_folder"] + types_docs[number_type]
            doc = docx.Document(path)
            for item in data:
                i = next(g)
                doc.tables[0].add_row()
                doc.tables[0].rows[i].cells[0].text = str(i)
                doc.tables[0].rows[i].cells[1].text = " ".join(item[:3]) # ФИО
                doc.tables[0].rows[i].cells[2].text = item[3]   # профессия
                doc.tables[0].rows[i].cells[3].text = "Сдал №" + item[5]
            path = self.path["print_folder"] + types_docs[number_type]
            doc.save(path)
        except:
            return msg_er(self, GET_FILE + path)


class CountPeople(QDialog):
    def __init__(self, parent=None):
        self.status_ = True
        self.conf = Ini(self)
        ui_file = self.conf.get_path_ui("new_TB")
        if not ui_file:
            self.status_ = False
            return
        try:
            uic.loadUi(ui_file, self)
        except:
            self.status_ = False
            return
        super(CountPeople, self).__init__(parent)
        # my_pass
        self.parent = parent
        self.table = "workers"
        self.b_ok.clicked.connect(self.ev_ok)
        self.b_cancel.clicked.connect(self.ev_cancel)
        count_people = len(self.parent.db.get_data("*", self.table))
        if count_people == ERR:
            return
        self.count.setMaximum(count_people)
        print(count_people)

    def ev_ok(self):
        self.parent.count_people_tb = self.count.value()
        self.close()
        return

    def ev_cancel(self):
        self.parent.count_people_tb = -1
        self.close()



