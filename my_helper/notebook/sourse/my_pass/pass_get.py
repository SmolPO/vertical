import docx
from my_helper.notebook.sourse.database import *
from my_helper.notebook.sourse.my_pass.pass_template import TempPass


class GetPass(TempPass):
    def __init__(self, parent):
        self.status_ = True
        self.conf = Ini(self)
        ui_file = self.conf.get_path_ui("pass_get")
        if not ui_file:
            self.status_ = False
            return
        super(GetPass, self).__init__(ui_file, parent, "workers")
        if not self.status_:
            return
        # my_pass
        self.parent = parent

        self.d_from.setDate(dt.datetime.now().date())
        self.d_to.setDate(from_str(".".join([str(count_days[dt.datetime.now().month - 1]),
                                             str(dt.datetime.now().month),
                                             str(dt.datetime.now().year)])))
        self.list_ui = (self.worker_1, self.worker_2, self.worker_3, self.worker_4, self.worker_5,
                        self.worker_6, self.worker_7, self.worker_8, self.worker_9, self.worker_10)
        self.data = {"customer": "", "company": "", "start_date": "", "end_date": "",
                     "contract": "", "date_contract": "", "number": "", "date": ""}
        self.main_file += "/pass_get.docx"
        self.init_workers()
        self.init_contracts()

    def init_contracts(self):
        contracts = self.parent.db.get_data("id, name", "contracts")
        if contracts == ERR or not contracts:
            return ERR
        for row in contracts:
            self.cb_contract.addItem(row[0] + "." + row[1])

    def init_workers(self):
        for item in self.list_ui:
            item.addItem("(нет)")
            item.activated[str].connect(self.new_worker)
            item.setEnabled(False)
        self.list_ui[0].setEnabled(True)
        people = self.parent.db.get_data("family, name, surname, post, passport, "
                                         "passport_got, birthday, adr,  live_adr", "workers")
        if people == ERR or not people:
            return ERR
        for name in people:
            family = name[0] + " " + ".".join([name[1][0], name[2][0]]) + "."
            for item in self.list_ui:
                item.addItem(family)

    def _get_data(self):
        self.data["start_date"] = self.d_from.text()
        self.data["end_date"] = self.d_to.text()
        self.data["customer"] = self.parent.customer
        self.data["company"] = self.parent.company
        rows = self.parent.db.get_data("id, number, date", "contracts")
        if rows == ERR:
            return ERR
        for contract in rows:
            if self.cb_contract.currentText().split(".")[0] == str(contract[0]):
                self.data["contract"] = contract[1]
                self.data["date_contract"] = contract[2]
        return True

    # обработчики кнопок
    def _ev_ok(self):
        if self.list_ui[0].currentText() == NOT:
            msg_er(self, ADD_PEOPLE)
            return False
        if self.cb_contract.currentText() == NOT:
            msg_info(self, ADD_CONTRACT)
            return False
        return True

    def _create_data(self, path):
        # Заполнить таблицу
        workers = []
        for elem in self.list_ui:
            if elem.currentText() != NOT:
                workers.append(self.get_worker(elem.currentText()))
        i = 1
        try:
            doc = docx.Document(path)
            for people in workers:
                doc.tables[1].add_row()
                doc.tables[1].rows[i].cells[0].text = str(i)
                doc.tables[1].rows[i].cells[1].text = " ".join(people[0:3])
                doc.tables[1].rows[i].cells[2].text = people[3]
                doc.tables[1].rows[i].cells[3].text = people[6]
                doc.tables[1].rows[i].cells[4].text = " ".join(people[4:6])
                doc.tables[1].rows[i].cells[5].text = people[7]
                doc.tables[1].rows[i].cells[6].text = people[8]
                i += 1
            doc.save(path)
        except:
            return msg_er(self, GET_FILE)

    def check_input(self):
        return True
