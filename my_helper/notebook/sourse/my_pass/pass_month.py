from PyQt5.QtCore import Qt
from PyQt5.QtCore import QDate as Date
from PyQt5.QtWidgets import QMessageBox as mes
import datetime as dt
from configparser import ConfigParser
import docxtpl
import os
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from my_helper.notebook.sourse.my_pass.pass_template import TempPass
from my_helper.notebook.sourse.database import *
import logging
import docx
#  logging.basicConfig(filename=get_path("path") + "/log_file.log", level=logging.INFO)
designer_file = get_path_ui("pass_month")


class MonthPass(TempPass):
    def __init__(self, parent):
        super(MonthPass, self).__init__(designer_file, parent, "workers")
        if not self.status_:
            return
        self.b_save.clicked.connect(self.save_pattern)
        self.b_kill.clicked.connect(self.kill_pattern)
        self.b_open.clicked.connect(self.my_open_file)
        self.count_people = 0
        self.d_from.setDate(dt.datetime.now().date())
        self.d_to.setDate(from_str(".".join([str(count_days[dt.datetime.now().month - 1]),
                                                   str(dt.datetime.now().month),
                                                   str(dt.datetime.now().year)])))
        self.cb_all.stateChanged.connect(self.set_enabled_workers)
        self.cb_manual_set.stateChanged.connect(self.set_dates)

        self.list_ui = (self.worker_1, self.worker_2, self.worker_3, self.worker_4, self.worker_5, self.worker_6,
                        self.worker_7, self.worker_8, self.worker_9, self.worker_10)
        self.data = {"customer": "", "company": "", "start_date": "", "end_date": "", "number": "", "date": ""}
        self.init_workers()
        self.init_cb_month()
        self.set_dates(self.cb_manual_set.isChecked())
        self.vac_path = self.main_file + "/Вакцинация_2.docx"
        self.main_file += "/pass_month.docx"


    # инициализация
    def init_cb_month(self):
        month = dt.datetime.now().month - 1
        for elem in self.list_month[month:]:
            self.cb_month.addItem(elem)

    def init_workers(self):
        for item in self.list_ui:
            item.addItem("(нет)")
            item.activated[str].connect(self.new_worker)
            item.setEnabled(False)
        self.list_ui[0].setEnabled(True)
        for people in self.all_people:
            family = str(people[-1]) + ". " + short_name(people)
            if family[-2] != 3:
                for item in self.list_ui:
                    item.addItem(family)

    # флаг на выбор всех
    def set_enabled_workers(self, state):
        for elem in self.list_ui:
            elem.setEnabled(state != Qt.Checked)
        if state != Qt.Checked:
            self.new_worker()

    # получить данные
    # для заполнения текста
    def _get_data(self):
        if self.cb_manual_set.isChecked():
            self.data["start_date"] = self.d_from.text()
            self.data["end_date"] = self.d_to.text()
            return
        next_month = self.list_month.index(self.cb_month.currentText()) + 1
        # если конец года: увеличить год и месяц в 1
        if next_month == 13:
            try:
                self.new_year__week = get_from_ini('config', 'new_year')
            except:
                return msg(self, my_errors["3_get_db"])
            next_day = self.new_year__week
            next_month = "01"  # MessageBox для ввода первого дня
            next_year = str(dt.datetime.now().year + 1)
        else:
            next_day = "01"
            next_month = str(next_month)
            if int(next_month) < 10:
                next_month = "0" + next_month
            next_year = str(dt.datetime.now().year)
        if int(next_year) / 4 == 0:
            end_next_month = str(count_days[12])
        else:
            end_next_month = str(count_days[int(next_month) - 1])
        self.data["start_date"] = ".".join((next_day, next_month, next_year))
        self.data["end_date"] = ".".join((end_next_month, next_month, next_year))
     
    # обработчики кнопок
    def _create_data(self, path):
        # Заполнить таблицу
        people = []
        doc = docx.Document(path)
        if self.cb_all.isChecked():
            people = self.all_people
        else:
            for elem in self.list_ui:
                if elem.currentText() != "(нет)":
                    people.append(self.check_row(elem.currentText()))
        i = 1
        people.sort(key=lambda x: x[0])
        for item in people:
            doc.tables[1].add_row()
            doc.tables[1].rows[i].cells[0].text = str(i)
            doc.tables[1].rows[i].cells[1].text = " ".join(item[0:3])
            doc.tables[1].rows[i].cells[2].text = item[3]
            doc.tables[1].rows[i].cells[3].text = item[6]
            doc.tables[1].rows[i].cells[4].text = " ".join(item[4:6])
            doc.tables[1].rows[i].cells[5].text = item[7]
            doc.tables[1].rows[i].cells[6].text = item[8]
            i += 1
        doc.save(path)
        os.startfile(path)
        set_next_number(self.number.value() + 1)
        self.create_vac(people)

    def check_input(self):
        if self.list_ui[0].currentText() == "(нет)" and not self.cb_all.isChecked():
            return msg(self, my_errors["14_add_people"])
        return True

    def _ev_ok(self):
        return True

    def create_vac(self, people):
        people.sort(key=lambda x: x[-2])
        list_people = {"2 дозы": [], "1 доза": [], "болел": []}
        for item in people:
            list_people[item[-2]].append(item)
        doc = docx.Document(self.vac_path)
        next_id = get_next_number()
        doc.tables[0].rows[0].cells[0].text = "Исх. " + str(next_id)
        doc.tables[0].rows[1].cells[0].text = "от " + self.d_note.text()
        ind = 0
        for key in list_people.keys():
            list_people[key].sort(key=lambda x: x[0][0])
            flag = True
            for man in list_people[key]:
                if flag:
                    if not man:
                        continue
                    ind = 1
                    if key == "2 дозы":
                        note = "2 дозы"
                        column = 6
                        self.init_table(doc, note, column)
                        self.init_SP5(doc, ind)
                    elif key == "1 доза":
                        note = "1 доза"
                        column = 5
                        self.init_table(doc, note, column)
                        self.init_LT(doc, ind)
                    elif key == "болел":
                        note = "болел"
                        column = 5
                        self.init_table(doc, note, column)
                        self.init_LT(doc, ind)
                    flag = False
                doc.tables[-1].add_row()
                add_list = [str(ind), short_name(man[:3]), man[3]]
                if man[-2] == "2 дозы":
                    add_list += [man[-6], man[-5], man[-4]]
                elif man[-2] == "1 дозы":
                    add_list += [man[-6], man[-4]]
                elif man[-2] == "болел":
                    add_list += [man[-6], man[-3]]
                for i in range(len(add_list)):
                    doc.tables[-1].rows[1].cells[i].text = add_list[i]
                ind += 1
        self.add_footer(doc)
        path = get_path("path") + get_path("path_notes_docs") + "/vac_" + str(dt.datetime.now().date()) + ".docx"
        doc.save(path)
        os.startfile(path)
        pass

    def init_SP5(self, doc, ind):
        doc.tables[-1].rows[0].cells[3].text = "Дата первой прививки"
        doc.tables[-1].rows[0].cells[4].text = "Дата второй прививки"
        doc.tables[-1].rows[0].cells[5].text = "Место вакцинации"
        pass

    def init_LT(self, doc, ind):
        doc.tables[-1].rows[0].cells[3].text = "Дата прививки"
        doc.tables[-1].rows[0].cells[4].text = "Место вакцинации"
        pass

    def init_CV(self, doc,  ind):
        doc.tables[-1].rows[0].cells[3].text = "Номер сертификата"
        doc.tables[-1].rows[0].cells[4].text = "Дата получения"
        pass

    def init_table(self, doc, note, column):
        doc.add_paragraph(note)
        doc.add_table(rows=1, cols=column)
        doc.tables[-1].rows[0].cells[0].text = "№"
        doc.tables[-1].rows[0].cells[1].text = "ФИО"
        doc.tables[-1].rows[0].cells[2].text = "Должность"

    def add_footer(self, doc):
        boss = self.parent.db.execute("SELECT big_post, big_boss "
                                      "FROM company "
                                      "WHERE company = '{0}'".format(self.parent.company))
        big_boss = self.parent.db.cursor.fetchall()
        note = big_boss[0][0] + "_"*15 + short_name(big_boss[0][1].split(" "))
        doc.add_paragraph(note)

    def init_part(self, key):
        return False