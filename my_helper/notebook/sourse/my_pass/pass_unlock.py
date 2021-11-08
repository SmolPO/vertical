from PyQt5.QtWidgets import QMessageBox as mes
import datetime as dt
import docx
import os
import logging
import openpyxl
import pymorphy2
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from my_helper.notebook.sourse.my_pass.pass_template import TempPass
from my_helper.notebook.sourse.database import *


class UnlockPass(TempPass):
    def __init__(self, parent):
        self.status_ = True
        self.conf = Ini(self)
        ui_file = self.conf.get_path_ui("pass_unlock")
        if not ui_file:
            self.status_ = False
            return
        super(UnlockPass, self).__init__(ui_file, parent, "workers")
        if not self.status_:
            return
        self.parent = parent
        # my_pass
        self.d_from.setDate(dt.datetime.now().date())
        self.d_to.setDate(from_str(".".join([str(count_days[dt.datetime.now().month - 1]),
                                             str(dt.datetime.now().month),
                                             str(dt.datetime.now().year)])))
        self.cb_all_days.stateChanged.connect(self.all_days)
        self.rows_from_db = self.parent.db.get_data("*", self.table)
        if self.rows_from_db == ERR:
            self.status_ = False
            return
        self.init_workers()
        self.data = {"number": "", "data": "", "customer": "", "company": "", "start_date": "", "end_date": "",
                     "post": "", "family": "", "name": "", "surname": "", "adr": ""}
        self.vac_path = self.main_file + "/Вакцинация_1.docx"
        self.main_file += "/Разблокировка.docx"
        self.count_days = 14

    def init_workers(self):
        for people in self.all_people:
            if people[-2] != 3:
                self.cb_worker.addItem(str(people[-1]) + ". " + short_name(people))

    def _get_data(self):
        family = self.cb_worker.currentText()
        morph = pymorphy2.MorphAnalyzer()
        people = self.check_row(family)  # на форме фамилия в виде Фамилия И.
        post = people[3]
        if post in dictionary.keys():
            post = dictionary[post]['datv']
        else:
            post = morph.parse(people[3])[0].inflect({'datv'})[0]
        self.data["family"] = morph.parse(people[0])[0].inflect({'datv'})[0].capitalize()
        self.data["name"] = morph.parse(people[1])[0].inflect({'datv'})[0].capitalize()
        self.data["surname"] = morph.parse(people[2])[0].inflect({'datv'})[0].capitalize()
        self.data["post"] = post
        self.data["adr"] = people[8]
        self.data["start_date"] = self.d_from.text()
        self.data["end_date"] = self.d_to.text()
        self.count_days = self.sb_days.value()


    def check_input(self):
        return True

    def _ev_ok(self):
        return True

    def all_days(self, state):
        self.sb_days.setEnabled(not state)
        self.sb_days.setValue(14)
        self.count_days = 14
        pass

    def _create_data(self, _):
        family = self.cb_worker.currentText()
        if self.create_vac(family) == ERR:
            return ERR
        pass

    def create_vac(self, family):
        note = ["Настоящим письмом информируем Вас о прохождение вакцинации от Covid-19 сотрудником ООО «Вертикаль»"]
        people = self.check_row(family)
        data_vac = people[-6:-1]
        try:
            doc = docx.Document(self.vac_path)
        except:
            return ERR
        next_id = self.conf.set_next_number(int(self.number.value()) + 1)
        doc.tables[0].rows[0].cells[0].text = "Исх. " + str(next_id)
        doc.tables[0].rows[1].cells[0].text = "от " + self.d_note.text()
        doc.tables[1].rows[1].cells[0].text = "1"
        doc.tables[1].rows[1].cells[1].text = " ".join(people[:3])
        doc.tables[1].rows[1].cells[2].text = people[3]
        p = CT_P.add_p_before(doc.tables[1]._element)
        p2 = Paragraph(p, doc.tables[1]._parent)
        p2.text = note
        if data_vac[4] == SPUTNIK:
            doc.tables[1].add_column(200)
            doc.tables[1].rows[0].cells[3].text = "Дата первой прививки"
            doc.tables[1].rows[0].cells[4].text = "Дата второй прививки"
            doc.tables[1].rows[0].cells[5].text = "Место вакцинации"

            doc.tables[1].rows[1].cells[3].text = data_vac[0]
            doc.tables[1].rows[1].cells[4].text = data_vac[1]
            doc.tables[1].rows[1].cells[5].text = data_vac[2]

        elif data_vac[4] == SP_LITE:
            self.data["d_vac_1"] = data_vac[0]
            self.data["place"] = data_vac[2]
            doc.tables[1].rows[0].cells[3].text = "Дата прививки"
            doc.tables[1].rows[0].cells[4].text = "Место вакцинации"
            doc.tables[1].rows[1].cells[3].text = data_vac[0]
            doc.tables[1].rows[1].cells[4].text = data_vac[2]

        elif data_vac[4] == COVID:
            self.data["d_vac_1"] = data_vac[0]
            self.data["vac_doc"] = data_vac[3][2:]
            doc.tables[1].rows[0].cells[3].text = "Номер сертификата"
            doc.tables[1].rows[0].cells[4].text = "Дата получения"

            doc.tables[1].rows[1].cells[3].text = data_vac[3]
            doc.tables[1].rows[1].cells[4].text = data_vac[0]
        next_id = self.conf.set_next_number(int(self.number.value()) + 1)
        path_1 = self.conf.get_path("path")
        path_2 = self.conf.get_path("path_notes_docs")
        if path_1 == ERR or path_2 == ERR:
            return ERR
        path = path_1 + path_2 + "/" + str(next_id) + "_" + self.d_note.text() + ".docx"
        try:
            doc.save(path)
            os.startfile(path)
        except:
            return ERR
